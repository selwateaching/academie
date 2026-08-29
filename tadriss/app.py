import html
import json
import os
import re
from datetime import datetime
from functools import wraps
from pathlib import Path

from dotenv import load_dotenv
from flask import (
    Flask,
    abort,
    jsonify,
    redirect,
    render_template,
    request,
    send_from_directory,
    session,
    url_for,
)
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_login import LoginManager, current_user, login_required, login_user, logout_user
from sqlalchemy import inspect, text

from ai import generate_document
from models import (
    DEFAULT_STATE,
    SUBSCRIPTION_DAYS,
    TRIAL_DAYS,
    Classe,
    Eleve,
    Evaluation,
    Note,
    Presence,
    User,
    UserState,
    db,
    generate_access_code,
    mention_for,
    moyenne_eleve,
)

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
GENERATED_DIR = BASE_DIR / "generated"
GENERATED_DIR.mkdir(exist_ok=True)

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret-change-me")
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get(
    "DATABASE_URL", f"sqlite:///{BASE_DIR / 'tadriss.db'}"
)
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["BANK_HOLDER"] = os.environ.get("BANK_HOLDER", "")
app.config["BANK_NAME"] = os.environ.get("BANK_NAME", "")
app.config["BANK_RIB"] = os.environ.get("BANK_RIB", "")

db.init_app(app)

login_manager = LoginManager(app)
login_manager.login_view = "login"

limiter = Limiter(get_remote_address, app=app, storage_uri="memory://", default_limits=[])

ADMIN_EMAIL = os.environ.get("ADMIN_EMAIL", "").strip().lower()


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def admin_required(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            abort(403)
        return fn(*args, **kwargs)

    return wrapper


def subscriber_required(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if not current_user.is_active_subscriber:
            abort(403)
        return fn(*args, **kwargs)

    return wrapper


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------


@app.errorhandler(429)
def too_many_requests(_exc):
    return (
        "Trop de tentatives. Veuillez patienter une minute avant de réessayer.",
        429,
        {"Content-Type": "text/plain; charset=utf-8"},
    )


@app.route("/signup", methods=["GET", "POST"])
@limiter.limit("10 per minute", methods=["POST"])
def signup():
    if current_user.is_authenticated:
        return redirect(url_for("index"))
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        name = request.form.get("name", "").strip()
        if not email or "@" not in email or len(password) < 6:
            return render_template(
                "signup.html",
                error="Email valide et mot de passe d'au moins 6 caractères requis.",
            )
        if User.query.filter_by(email=email).first():
            return render_template(
                "signup.html", error="Un compte existe déjà avec cet email."
            )
        user = User(email=email, name=name, is_admin=(email == ADMIN_EMAIL))
        user.set_password(password)
        db.session.add(user)
        db.session.flush()
        db.session.add(
            UserState(user_id=user.id, data=json.dumps(DEFAULT_STATE, ensure_ascii=False))
        )
        seed_demo_classe(user)
        db.session.commit()
        login_user(user)
        return redirect(url_for("index"))
    return render_template("signup.html")


def seed_demo_classe(user):
    """Crée une classe et quelques élèves de démonstration pour découvrir l'appli."""
    classe = Classe(
        user_id=user.id, nom="4AM B", matiere="Mathématiques", annee_scolaire=DEFAULT_STATE["year"]
    )
    db.session.add(classe)
    db.session.flush()
    demo_eleves = [
        ("Amine", "Belaïd"),
        ("Lina", "Kaci"),
        ("Yanis", "Meziane"),
        ("Sara", "Amrani"),
    ]
    for prenom, nom in demo_eleves:
        db.session.add(Eleve(classe_id=classe.id, prenom=prenom, nom=nom))


@app.route("/login", methods=["GET", "POST"])
@limiter.limit("10 per minute", methods=["POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("index"))
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        user = User.query.filter_by(email=email).first()
        if user and user.check_password(password):
            login_user(user, remember=True)
            return redirect(url_for("index"))
        return render_template("login.html", error="Email ou mot de passe incorrect.")
    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("login"))


@app.route("/billing")
@login_required
def billing():
    return render_template("billing.html", user=current_user)


# ---------------------------------------------------------------------------
# App shell
# ---------------------------------------------------------------------------


@app.route("/")
@login_required
def index():
    if not current_user.is_active_subscriber:
        return redirect(url_for("billing"))
    return render_template("index.html", user=current_user)


# ---------------------------------------------------------------------------
# API : état de l'abonné (remplace le localStorage du navigateur)
# ---------------------------------------------------------------------------


@app.route("/api/profile", methods=["PUT"])
@login_required
def api_profile_update():
    data = request.get_json(force=True, silent=True) or {}
    if "name" in data:
        current_user.name = (data["name"] or "").strip()
    if "phone" in data:
        current_user.phone = (data["phone"] or "").strip()
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/state", methods=["GET"])
@login_required
def get_state():
    state = current_user.state
    if not state:
        return jsonify(DEFAULT_STATE)
    try:
        return jsonify(json.loads(state.data))
    except (TypeError, ValueError):
        return jsonify(DEFAULT_STATE)


@app.route("/api/state", methods=["POST"])
@login_required
@subscriber_required
def save_state():
    payload = request.get_json(force=True, silent=True)
    if not isinstance(payload, dict):
        abort(400)
    state = current_user.state
    if not state:
        state = UserState(user_id=current_user.id)
        db.session.add(state)
    state.data = json.dumps(payload, ensure_ascii=False)
    state.updated_at = datetime.utcnow()
    db.session.commit()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# API : génération IA
# ---------------------------------------------------------------------------


@app.route("/api/ai/generate", methods=["POST"])
@login_required
@subscriber_required
def api_ai_generate():
    payload = request.get_json(force=True, silent=True) or {}
    document = generate_document(
        payload.get("type", "sheet"),
        payload.get("level", ""),
        payload.get("subject", ""),
        payload.get("lang", "auto"),
        payload.get("prompt", ""),
    )
    return jsonify({"ok": True, "document": document})


# ---------------------------------------------------------------------------
# API : classes, élèves, notes et présences (appel)
# ---------------------------------------------------------------------------


def owned_classe_or_404(classe_id):
    classe = Classe.query.get_or_404(classe_id)
    if classe.user_id != current_user.id:
        abort(404)
    return classe


def owned_eleve_or_404(eleve_id):
    eleve = Eleve.query.get_or_404(eleve_id)
    if eleve.classe.user_id != current_user.id:
        abort(404)
    return eleve


def owned_evaluation_or_404(evaluation_id):
    evaluation = Evaluation.query.get_or_404(evaluation_id)
    if evaluation.classe.user_id != current_user.id:
        abort(404)
    return evaluation


def serialize_eleve(eleve):
    moyenne = moyenne_eleve(eleve)
    return {
        "id": eleve.id,
        "classe_id": eleve.classe_id,
        "prenom": eleve.prenom,
        "nom": eleve.nom,
        "date_naissance": eleve.date_naissance,
        "responsable_nom": eleve.responsable_nom,
        "responsable_lien": eleve.responsable_lien,
        "responsable_tel": eleve.responsable_tel,
        "responsable_email": eleve.responsable_email,
        "observations": eleve.observations,
        "access_code": eleve.access_code,
        "moyenne": moyenne,
        "mention": mention_for(moyenne),
    }


def save_notes_bulk(evaluation, notes_dict):
    valid_ids = {e.id for e in evaluation.classe.eleves}
    for eleve_id_str, valeur in (notes_dict or {}).items():
        try:
            eleve_id = int(eleve_id_str)
        except (TypeError, ValueError):
            continue
        if eleve_id not in valid_ids:
            continue
        note = Note.query.filter_by(evaluation_id=evaluation.id, eleve_id=eleve_id).first()
        if valeur in (None, ""):
            if note:
                db.session.delete(note)
            continue
        try:
            v = float(valeur)
        except (TypeError, ValueError):
            continue
        if note:
            note.valeur = v
        else:
            db.session.add(Note(evaluation_id=evaluation.id, eleve_id=eleve_id, valeur=v))
    db.session.commit()


@app.route("/api/classes", methods=["GET"])
@login_required
def api_classes_list():
    classes = Classe.query.filter_by(user_id=current_user.id).order_by(Classe.nom).all()
    return jsonify(
        [
            {
                "id": c.id,
                "nom": c.nom,
                "matiere": c.matiere,
                "annee_scolaire": c.annee_scolaire,
                "nb_eleves": len(c.eleves),
            }
            for c in classes
        ]
    )


@app.route("/api/classes", methods=["POST"])
@login_required
@subscriber_required
def api_classes_create():
    data = request.get_json(force=True, silent=True) or {}
    nom = (data.get("nom") or "").strip()
    if not nom:
        abort(400)
    classe = Classe(
        user_id=current_user.id,
        nom=nom,
        matiere=(data.get("matiere") or "").strip(),
        annee_scolaire=(data.get("annee_scolaire") or "").strip() or DEFAULT_STATE["year"],
    )
    db.session.add(classe)
    db.session.commit()
    return jsonify(
        {
            "id": classe.id,
            "nom": classe.nom,
            "matiere": classe.matiere,
            "annee_scolaire": classe.annee_scolaire,
        }
    )


@app.route("/api/classes/<int:classe_id>", methods=["PUT"])
@login_required
@subscriber_required
def api_classes_update(classe_id):
    classe = owned_classe_or_404(classe_id)
    data = request.get_json(force=True, silent=True) or {}
    nom = (data.get("nom") or "").strip()
    if nom:
        classe.nom = nom
    classe.matiere = (data.get("matiere") or "").strip()
    if "annee_scolaire" in data:
        classe.annee_scolaire = (data["annee_scolaire"] or "").strip() or classe.annee_scolaire
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/classes/<int:classe_id>", methods=["DELETE"])
@login_required
@subscriber_required
def api_classes_delete(classe_id):
    classe = owned_classe_or_404(classe_id)
    db.session.delete(classe)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/classes/<int:classe_id>/eleves", methods=["GET"])
@login_required
def api_eleves_list(classe_id):
    classe = owned_classe_or_404(classe_id)
    return jsonify([serialize_eleve(e) for e in classe.eleves])


@app.route("/api/classes/<int:classe_id>/eleves", methods=["POST"])
@login_required
@subscriber_required
def api_eleves_create(classe_id):
    classe = owned_classe_or_404(classe_id)
    data = request.get_json(force=True, silent=True) or {}
    prenom = (data.get("prenom") or "").strip()
    nom = (data.get("nom") or "").strip()
    if not prenom or not nom:
        abort(400)
    eleve = Eleve(
        classe_id=classe.id,
        prenom=prenom,
        nom=nom,
        date_naissance=(data.get("date_naissance") or "").strip(),
        responsable_nom=(data.get("responsable_nom") or "").strip(),
        responsable_lien=(data.get("responsable_lien") or "").strip(),
        responsable_tel=(data.get("responsable_tel") or "").strip(),
        responsable_email=(data.get("responsable_email") or "").strip(),
        observations=(data.get("observations") or "").strip(),
    )
    db.session.add(eleve)
    db.session.commit()
    return jsonify(serialize_eleve(eleve))


@app.route("/api/eleves/<int:eleve_id>", methods=["PUT"])
@login_required
@subscriber_required
def api_eleve_update(eleve_id):
    eleve = owned_eleve_or_404(eleve_id)
    data = request.get_json(force=True, silent=True) or {}
    for field in (
        "prenom",
        "nom",
        "date_naissance",
        "responsable_nom",
        "responsable_lien",
        "responsable_tel",
        "responsable_email",
        "observations",
    ):
        if field in data:
            setattr(eleve, field, (data[field] or "").strip())
    db.session.commit()
    return jsonify(serialize_eleve(eleve))


@app.route("/api/eleves/<int:eleve_id>", methods=["DELETE"])
@login_required
@subscriber_required
def api_eleve_delete(eleve_id):
    eleve = owned_eleve_or_404(eleve_id)
    db.session.delete(eleve)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/eleves/<int:eleve_id>/regenerate_code", methods=["POST"])
@login_required
@subscriber_required
def api_eleve_regenerate_code(eleve_id):
    eleve = owned_eleve_or_404(eleve_id)
    eleve.access_code = generate_access_code()
    db.session.commit()
    return jsonify({"access_code": eleve.access_code})


@app.route("/api/eleves/<int:eleve_id>/bulletin", methods=["GET"])
@login_required
def api_eleve_bulletin(eleve_id):
    eleve = owned_eleve_or_404(eleve_id)
    evaluations = Evaluation.query.filter_by(classe_id=eleve.classe_id).order_by(Evaluation.date).all()
    rows = []
    for ev in evaluations:
        note = Note.query.filter_by(evaluation_id=ev.id, eleve_id=eleve.id).first()
        rows.append(
            {
                "intitule": ev.intitule,
                "type": ev.type,
                "date": ev.date,
                "coefficient": ev.coefficient,
                "valeur": note.valeur if note else None,
            }
        )
    moyenne = moyenne_eleve(eleve)
    return jsonify(
        {
            "eleve": serialize_eleve(eleve),
            "classe": eleve.classe.nom,
            "evaluations": rows,
            "moyenne": moyenne,
            "mention": mention_for(moyenne),
        }
    )


@app.route("/api/eleves/<int:eleve_id>/absences", methods=["GET"])
@login_required
def api_eleve_absences(eleve_id):
    eleve = owned_eleve_or_404(eleve_id)
    records = (
        Presence.query.filter(Presence.eleve_id == eleve.id, Presence.statut.in_(["A", "R"]))
        .order_by(Presence.date.desc())
        .all()
    )
    return jsonify([{"date": r.date, "statut": r.statut} for r in records])


@app.route("/api/classes/<int:classe_id>/evaluations", methods=["GET"])
@login_required
def api_evaluations_list(classe_id):
    classe = owned_classe_or_404(classe_id)
    evaluations = Evaluation.query.filter_by(classe_id=classe.id).order_by(Evaluation.date.desc()).all()
    out = []
    for ev in evaluations:
        notes = {n.eleve_id: n.valeur for n in ev.notes}
        out.append(
            {
                "id": ev.id,
                "intitule": ev.intitule,
                "type": ev.type,
                "date": ev.date,
                "coefficient": ev.coefficient,
                "notes": notes,
            }
        )
    return jsonify(out)


@app.route("/api/classes/<int:classe_id>/evaluations", methods=["POST"])
@login_required
@subscriber_required
def api_evaluations_create(classe_id):
    classe = owned_classe_or_404(classe_id)
    data = request.get_json(force=True, silent=True) or {}
    intitule = (data.get("intitule") or "").strip()
    if not intitule:
        abort(400)
    evaluation = Evaluation(
        classe_id=classe.id,
        intitule=intitule,
        type=(data.get("type") or "Contrôle").strip(),
        date=(data.get("date") or "").strip(),
        coefficient=int(data.get("coefficient") or 1),
    )
    db.session.add(evaluation)
    db.session.commit()
    save_notes_bulk(evaluation, data.get("notes") or {})
    return jsonify({"id": evaluation.id})


@app.route("/api/evaluations/<int:evaluation_id>", methods=["PUT"])
@login_required
@subscriber_required
def api_evaluation_update(evaluation_id):
    evaluation = owned_evaluation_or_404(evaluation_id)
    data = request.get_json(force=True, silent=True) or {}
    if data.get("intitule"):
        evaluation.intitule = data["intitule"].strip()
    if "type" in data:
        evaluation.type = (data["type"] or "").strip()
    if "date" in data:
        evaluation.date = (data["date"] or "").strip()
    if "coefficient" in data:
        evaluation.coefficient = int(data["coefficient"] or 1)
    db.session.commit()
    save_notes_bulk(evaluation, data.get("notes") or {})
    return jsonify({"ok": True})


@app.route("/api/evaluations/<int:evaluation_id>", methods=["DELETE"])
@login_required
@subscriber_required
def api_evaluation_delete(evaluation_id):
    evaluation = owned_evaluation_or_404(evaluation_id)
    db.session.delete(evaluation)
    db.session.commit()
    return jsonify({"ok": True})


@app.route("/api/classes/<int:classe_id>/presence", methods=["GET"])
@login_required
def api_presence_get(classe_id):
    classe = owned_classe_or_404(classe_id)
    date = request.args.get("date") or datetime.utcnow().strftime("%Y-%m-%d")
    records = {p.eleve_id: p.statut for p in Presence.query.filter_by(classe_id=classe.id, date=date).all()}
    return jsonify({"date": date, "presences": records})


@app.route("/api/classes/<int:classe_id>/presence", methods=["POST"])
@login_required
@subscriber_required
def api_presence_save(classe_id):
    classe = owned_classe_or_404(classe_id)
    data = request.get_json(force=True, silent=True) or {}
    date = (data.get("date") or "").strip()
    if not date:
        abort(400)
    valid_ids = {e.id for e in classe.eleves}
    for eleve_id_str, statut in (data.get("presences") or {}).items():
        try:
            eleve_id = int(eleve_id_str)
        except (TypeError, ValueError):
            continue
        if eleve_id not in valid_ids or statut not in ("P", "A", "R"):
            continue
        record = Presence.query.filter_by(eleve_id=eleve_id, date=date).first()
        if record:
            record.statut = statut
        else:
            db.session.add(Presence(classe_id=classe.id, eleve_id=eleve_id, date=date, statut=statut))
    db.session.commit()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Espace élève (public, accès par code — lecture seule)
# ---------------------------------------------------------------------------


def current_eleve():
    eleve_id = session.get("eleve_id")
    return Eleve.query.get(eleve_id) if eleve_id else None


@app.route("/eleve", methods=["GET"])
def eleve_login_page():
    if current_eleve():
        return redirect(url_for("eleve_espace"))
    return render_template("eleve_login.html", prefill=request.args.get("code", ""))


@app.route("/eleve/login", methods=["POST"])
def eleve_login():
    code = (request.form.get("code") or "").strip().replace(" ", "").lower()
    eleve = Eleve.query.filter_by(access_code=code).first() if code else None
    if not eleve:
        return render_template("eleve_login.html", error="Code invalide.", prefill="")
    session["eleve_id"] = eleve.id
    return redirect(url_for("eleve_espace"))


@app.route("/eleve/logout")
def eleve_logout():
    session.pop("eleve_id", None)
    return redirect(url_for("eleve_login_page"))


@app.route("/eleve/espace")
def eleve_espace():
    eleve = current_eleve()
    if not eleve:
        return redirect(url_for("eleve_login_page"))
    evaluations = Evaluation.query.filter_by(classe_id=eleve.classe_id).order_by(Evaluation.date).all()
    rows = []
    for ev in evaluations:
        note = Note.query.filter_by(evaluation_id=ev.id, eleve_id=eleve.id).first()
        rows.append(
            {
                "intitule": ev.intitule,
                "type": ev.type,
                "date": ev.date,
                "coefficient": ev.coefficient,
                "valeur": note.valeur if note else None,
            }
        )
    moyenne = moyenne_eleve(eleve)
    absences = (
        Presence.query.filter(Presence.eleve_id == eleve.id, Presence.statut.in_(["A", "R"]))
        .order_by(Presence.date.desc())
        .all()
    )
    return render_template(
        "eleve_espace.html",
        eleve=eleve,
        evaluations=rows,
        moyenne=moyenne,
        mention=mention_for(moyenne),
        absences=absences,
    )


# ---------------------------------------------------------------------------
# Export PDF / Word / HTML (par abonné)
# ---------------------------------------------------------------------------


def safe_name(s):
    s = re.sub(r"[^\w\-]+", "_", s or "", flags=re.UNICODE).strip("_")[:70]
    return s or "tadriss_document"


def make_html_doc(c):
    rtl = bool(c.get("rtl", False))
    direction = "rtl" if rtl else "ltr"
    lang = "ar" if rtl else "fr"
    sections = "".join(
        f"<section><h2>{html.escape(a)}</h2><p>{html.escape(b)}</p></section>"
        for a, b in c.get("sections", [])
    )
    title = html.escape(c.get("title", "Document pédagogique"))
    meta = html.escape(c.get("meta", ""))
    prompt_text = html.escape(c.get("prompt", ""))
    return f"""<!doctype html><html lang="{lang}" dir="{direction}"><head><meta charset="utf-8">
<title>{html.escape(c.get("title", "TADRISS"))}</title>
<style>@page{{size:A4;margin:18mm}}body{{font-family:"DejaVu Sans",Arial,sans-serif;color:#17203f;
line-height:1.7;font-size:11pt}}h1{{font-size:25pt;color:#101b4d;margin-bottom:4pt}}
.meta{{color:#75809d}}h2{{font-size:15pt;color:#101b4d;border-bottom:2px solid #ff8351;
padding-bottom:5pt;margin-top:20pt}}.rtl{{direction:rtl;text-align:right}}</style></head>
<body class="{'rtl' if rtl else ''}"><h1>{title}</h1><p class="meta">{meta}</p>
<p>{prompt_text}</p>{sections}<hr><small>Généré par TADRISS</small></body></html>"""


def make_docx_doc(c, path):
    from docx import Document
    from docx.shared import Pt

    d = Document()
    sec = d.sections[0]
    sec.top_margin = sec.bottom_margin = Pt(45)
    sec.left_margin = sec.right_margin = Pt(45)
    d.add_heading(c.get("title", "Document pédagogique"), 0)
    d.add_paragraph(c.get("meta", ""))
    d.add_paragraph(c.get("prompt", ""))
    for a, b in c.get("sections", []):
        d.add_heading(a, 1)
        d.add_paragraph(b)
    d.add_paragraph("Généré par TADRISS")
    d.save(path)


@app.route("/api/generate", methods=["POST"])
@login_required
@subscriber_required
def api_generate():
    c = request.get_json(force=True, silent=True) or {}
    user_dir = GENERATED_DIR / str(current_user.id)
    user_dir.mkdir(exist_ok=True)
    stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    base = safe_name(c.get("title", "document")) + "_" + stamp
    html_path = user_dir / f"{base}.html"
    pdf_path = user_dir / f"{base}.pdf"
    docx_path = user_dir / f"{base}.docx"
    try:
        from weasyprint import HTML

        html_content = make_html_doc(c)
        html_path.write_text(html_content, encoding="utf-8")
        HTML(string=html_content, base_url=str(BASE_DIR)).write_pdf(str(pdf_path))
        make_docx_doc(c, str(docx_path))
    except Exception:  # noqa: BLE001
        app.logger.exception("Échec de génération PDF/Word pour l'utilisateur %s", current_user.id)
        return (
            jsonify(
                {
                    "ok": False,
                    "error": "Impossible de générer le document pour le moment. Veuillez réessayer dans quelques instants.",
                }
            ),
            500,
        )
    return jsonify(
        {
            "ok": True,
            "files": {
                "html": url_for(
                    "generated_file", user_id=current_user.id, filename=html_path.name
                ),
                "pdf": url_for(
                    "generated_file", user_id=current_user.id, filename=pdf_path.name
                ),
                "docx": url_for(
                    "generated_file", user_id=current_user.id, filename=docx_path.name
                ),
            },
        }
    )


@app.route("/generated/<int:user_id>/<path:filename>")
@login_required
def generated_file(user_id, filename):
    if current_user.id != user_id and not current_user.is_admin:
        abort(403)
    return send_from_directory(GENERATED_DIR / str(user_id), filename)


# ---------------------------------------------------------------------------
# Admin : suivi des abonnés et paiements par virement
# ---------------------------------------------------------------------------


@app.route("/admin")
@login_required
@admin_required
def admin_panel():
    users = User.query.order_by(User.created_at.desc()).all()
    return render_template("admin.html", users=users, now=datetime.utcnow())


@app.route("/admin/mark_paid/<int:user_id>", methods=["POST"])
@login_required
@admin_required
def admin_mark_paid(user_id):
    user = User.query.get_or_404(user_id)
    days = int(request.form.get("days", SUBSCRIPTION_DAYS))
    user.mark_paid(days=days)
    db.session.commit()
    return redirect(url_for("admin_panel"))


@app.context_processor
def inject_globals():
    return {"trial_days": TRIAL_DAYS}


def ensure_column(table_name, column_name, ddl):
    """Ajoute une colonne manquante sur une table existante (mini-migration sans Alembic)."""
    inspector = inspect(db.engine)
    if table_name not in inspector.get_table_names():
        return
    existing = {c["name"] for c in inspector.get_columns(table_name)}
    if column_name not in existing:
        db.session.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {ddl}"))
        db.session.commit()


with app.app_context():
    db.create_all()
    ensure_column("user", "phone", "VARCHAR(40) DEFAULT ''")
    ensure_column("classe", "annee_scolaire", f"VARCHAR(20) DEFAULT '{DEFAULT_STATE['year']}'")
    if ADMIN_EMAIL:
        admin_user = User.query.filter_by(email=ADMIN_EMAIL).first()
        if admin_user and not admin_user.is_admin:
            admin_user.is_admin = True
            db.session.commit()


if __name__ == "__main__":
    app.run(
        host="0.0.0.0",
        port=int(os.environ.get("PORT", 8000)),
        debug=bool(os.environ.get("FLASK_DEBUG")),
    )
